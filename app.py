from flask import Flask, request, jsonify, Response
from openai import OpenAI
import os
import openai
from dotenv import load_dotenv
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
from io import StringIO
from pdfminer.high_level import extract_text
import docx
import re
from flask import request, jsonify
from collections import Counter
import PyPDF2
from flask import Flask, request, jsonify
# ✅ NLTK imports for keyword analysis
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import nltk

from flask import Flask, request, jsonify, render_template, send_file, make_response
import os
PDF_FOLDER = "generated_resumes"
os.makedirs(PDF_FOLDER, exist_ok=True)
from models import db
from auth import auth
from functools import wraps
from flask import session, redirect
from functools import wraps
from flask import session, redirect, url_for
from auth import auth as auth_bp
from flask import send_file
import joblib
from sklearn.metrics.pairwise import cosine_similarity

model = joblib.load("resume_model.pkl")
vectorizer = joblib.load("resume_vectorizer.pkl")





client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))



def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("auth.login"))
        return f(*args, **kwargs)
    return decorated_function



# Optional but needed for PDF generation if you're using WeasyPrint
try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

# Ensure required NLTK data is available
for resource in ["punkt", "punkt_tab", "stopwords"]:
    try:
        nltk.data.find(f"tokenizers/{resource}")
    except LookupError:
        nltk.download(resource)


# ✅ Initialize Flask App Only ONCE
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})  # Enable CORS globally


# ✅ Database Config (USE MYSQL ONLY)
app.config["SECRET_KEY"] = "a-very-secret-key-change-me"
app.config["SQLALCHEMY_DATABASE_URI"] = "mysql+pymysql://root:@localhost:3308/resume_db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

# ✅ Attach database
db.init_app(app)

# ✅ Register blueprint
app.register_blueprint(auth_bp)


# ✅ Load env
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# ✅ Define Resume model
class Resume(db.Model):
    __tablename__ = "resume"

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"))  # ✅ ADD THIS

    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100))
    phone = db.Column(db.String(20))
    address = db.Column(db.String(200))
    job_title = db.Column(db.String(100))
    summary = db.Column(db.Text)

    skills = db.Column(db.Text)
    languages = db.Column(db.Text)
    interests = db.Column(db.Text)
    references = db.Column(db.Text)
    education = db.Column(db.Text)
    experience = db.Column(db.Text)

    template = db.Column(db.String(50), default="template1")


# app.py (Ensure your PDF function is exactly this)

def extract_text_from_pdf(filepath):
    """Uses PyPDF2 to extract text from a PDF file."""
    if not PyPDF2:
        return "Error: PyPDF2 library is not installed."
    try:
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
    except Exception as e:
        # This catch is CRITICAL to prevent the server crash
        print(f"PDF EXTRACTION CRASH: {e}")
        return f"Error reading PDF: {str(e)}"


def extract_text_from_file(file_path):
    text = ""
    if file_path.endswith(".pdf"):
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text.strip()

# ------------- Parsing & Cleaning Helpers -------------
_BULLET_CHARS = r"[•●▪◦\-–—\u2022\u25CF\u25AA\u25E6\u2023\u2219✔✓✓]"

def _clean_text(text: str) -> str:
    if not text:
        return ""
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove non-printable characters but preserve newlines
    text = ''.join(ch for ch in text if ch.isprintable() or ch == '\n')
    # Replace many bullet characters with a normalized bullet marker
    text = re.sub(_BULLET_CHARS, "\n• ", text)
    # Remove repeated newlines (>2) to maintain block separation
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Trim each line and drop repeated identical lines
    lines = [ln.strip(' \t') for ln in text.splitlines()]
    cleaned = []
    prev = None
    for ln in lines:
        if ln == prev:
            continue
        cleaned.append(ln)
        prev = ln
    # Remove leading/trailing blank lines
    while cleaned and cleaned[0] == "":
        cleaned.pop(0)
    while cleaned and cleaned[-1] == "":
        cleaned.pop()
    return "\n".join(cleaned).strip()

def _split_items_by_separators(text: str):
    """Split skills / languages / interests into list by commas/semicolons/newlines/bullets."""
    if not text:
        return []
    t = _clean_text(text)
    parts = re.split(r"(?:\n+|,|;|\u2022|•)", t)
    out = []
    seen = set()
    for p in parts:
        s = p.strip(" \t\n\r.·•▪-")
        if len(s) < 2:
            continue
        key = s.lower()
        if key not in seen:
            seen.add(key)
            out.append(s)
    return out

_DATE_PATTERN = re.compile(r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec|\b\d{4}\b|\b\d{2}:\d{2}\b)", re.I)

def parse_experience_blocks(text: str):
    """
    Improved parser for Experience section.
    Produces structured blocks like:
    [
      {
        "title": "Construction Assistant Worker",
        "company": "Expeteers Pvt Ltd – Mutare, Zimbabwe",
        "dates": "January 2021 – December 2021",
        "duties": [
            "Assisted in site preparation and cleanup.",
            "Operated basic tools and helped with material installation.",
            ...
        ]
      },
      ...
    ]
    """
    if not text:
        return []

    txt = _clean_text(text)

    # Split experiences based on role patterns or year/date hints
    blocks = re.split(r'(?=\b[A-Z][a-z]+\s(?:Assistant|Worker|Technician|Engineer|Manager|Officer|Supervisor|Operator|Laborer|Mason|Foreman)\b)', txt)
    parsed = []

    for b in blocks:
        lines = [ln.strip() for ln in b.splitlines() if ln.strip()]
        if not lines:
            continue

        title, company, dates = "", "", ""
        duties = []

        # Title is usually the first line
        title = lines[0]

        # Detect company + location (often second line)
        if len(lines) > 1 and "–" in lines[1]:
            company = lines[1]
        elif len(lines) > 1:
            company = lines[1]

        # Detect dates (e.g., "January 2021 – December 2021")
        date_pattern = re.compile(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}\s*[–-]\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}", re.I)
        for ln in lines:
            if date_pattern.search(ln):
                dates = ln.strip()
                break

        # Extract duties (start with a bullet, number, or verb)
        for ln in lines[2:]:
            if re.match(r"^[•\-\*]", ln) or re.match(r"^(Assisted|Operated|Collaborated|Followed|Maintained|Mixed|Transported|Fixed|Reading|Smoothing)", ln):
                duties.append(ln.lstrip("•-* ").strip())

        parsed.append({
            "title": title,
            "company": company,
            "dates": dates,
            "duties": duties or [b]  # fallback
        })

    return parsed


def parse_education_blocks(text: str):
    """
    Parse education into list of blocks: { raw: '...', dates: '...' }
    """
    if not text:
        return []
    t = _clean_text(text)
    blocks = [b.strip() for b in re.split(r"\n\s*\n", t) if b.strip()]
    if len(blocks) == 1:
        # attempt to split by institution keywords
        candidates = re.split(r"(?=(?:High School|University|College|Institute|Academy|Association|Certificate|Diploma|BSc|BA|MSc|Degree))", t)
        blocks = [b.strip() for b in candidates if b.strip()]

    out = []
    for b in blocks:
        dates = " ".join(re.findall(r"\d{4}", b))
        out.append({"raw": b, "dates": dates})
    return out

def clean_resume_data(resume):
    """
    Convert resume model to cleaned, structured dict for templates.
    """
    return {
        "summary": _clean_text(getattr(resume, "summary", "") or ""),
        "skills": _split_items_by_separators(getattr(resume, "skills", "") or ""),
        "experience": parse_experience_blocks(getattr(resume, "experience", "") or ""),
        "education": parse_education_blocks(getattr(resume, "education", "") or ""),
        "languages": _split_items_by_separators(getattr(resume, "languages", "") or ""),
        "interests": _split_items_by_separators(getattr(resume, "interests", "") or ""),
        "references": _clean_text(getattr(resume, "references", "") or "")
    }
# Map template IDs to pre-made PDF files
TEMPLATE_FILES = {
    "template1": "static/resumes/template1.pdf",
    "template2": "static/resumes/template2.pdf",
    "template3": "static/resumes/template3.pdf",
    "template4": "static/resumes/template4.pdf"
}

@app.route("/resume/<template_id>")
def get_resume(template_id):
    file_path = TEMPLATE_FILES.get(template_id)
    if not file_path:
        return "Template not found", 404
    # Open PDF in browser instead of download
    return send_file(file_path, as_attachment=False)

@app.route("/favicon.ico")
def favicon():
    from flask import send_from_directory
    import os
    return send_from_directory(os.path.join(app.root_path, 'static'),
                               'favicon.ico', mimetype='image/vnd.microsoft.icon')

@app.route("/resume_tool")
def resume_tool():
    return render_template("resume_tool.html")



# ✅ Resume Dashboard Route
@app.route("/dashboard")
@login_required
def dashboard():
 return render_template("dashboard.html")

@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile():
    from models import User
    from werkzeug.security import check_password_hash, generate_password_hash

    user = User.query.get(session["user_id"])
    message = ""

    if request.method == "POST":
        old_pass = request.form.get("old_password")
        new_pass = request.form.get("new_password")
        confirm = request.form.get("confirm_password")

        if not check_password_hash(user.password, old_pass):
            message = "❌ Current password is incorrect"
        elif new_pass != confirm:
            message = "❌ Passwords do not match"
        elif len(new_pass) < 6:
            message = "❌ Password too short"
        else:
            user.password = generate_password_hash(new_pass)
            db.session.commit()
            message = "✅ Password updated successfully"

    return render_template("profile.html", user=user, message=message)

@app.route("/my-resumes")
@login_required
def resume_history():
    user_id = session["user_id"]

    resumes = Resume.query.filter_by(user_id=user_id).order_by(Resume.id.desc()).all()

    return render_template("resume_history.html", resumes=resumes)

@app.route("/delete_resume/<int:resume_id>")
@login_required
def delete_resume(resume_id):

    resume = Resume.query.filter_by(
        id=resume_id,
        user_id=session["user_id"]
    ).first_or_404()

    db.session.delete(resume)
    db.session.commit()

    return redirect("/my-resumes")



@app.route("/build_resume")
@login_required
def build_resume():
 return render_template("resume_builder_form.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("auth.login"))


# ✅ Resume Submission Route
@login_required
@app.route('/submit_resume', methods=['POST'])
def submit_resume():
    data = request.get_json()
    print("📩 Received resume data:", data)

    if not data:
        return jsonify({"error": "No JSON data received"}), 400

    import json
    from flask import session

    # Normalize lists -> JSON strings
    experience_json = json.dumps(data.get("experience", []))
    education_json = json.dumps(data.get("education", []))
    references_json = json.dumps(data.get("references", []))

    skills_json = json.dumps(data.get("skills", []))
    languages_json = json.dumps(data.get("languages", []))
    interests_json = json.dumps(data.get("interests", []))

    # DEFAULT template if not provided
    template_value = data.get("template") or "template1"

    # If resume_id provided -> update existing resume
    resume_id = data.get("resume_id")
    if resume_id:
        resume = Resume.query.filter_by(id=resume_id, user_id=session["user_id"]).first()
        if not resume:
            return jsonify({"error": "Resume not found"}), 404

        resume.name = data.get('name')
        resume.email = data.get('email')
        resume.phone = data.get('phone')
        resume.address = data.get('address')
        resume.job_title = data.get('job_title')
        resume.summary = data.get('summary')

        resume.experience = experience_json
        resume.education = education_json
        resume.references = references_json

        resume.skills = skills_json
        resume.languages = languages_json
        resume.interests = interests_json

        # store the template (overwrite only if provided)
        if data.get("template") is not None:
            resume.template = template_value

        db.session.commit()

        print(f"✅ Updated resume ID {resume.id} with template {resume.template}")
        return jsonify({"message": "Resume updated", "resume_id": resume.id, "template": resume.template}), 200

    # Otherwise create new resume
    resume = Resume(
        user_id=session["user_id"],
        name=data.get('name'),
        email=data.get('email'),
        phone=data.get('phone'),
        address=data.get('address'),
        job_title=data.get('job_title'),
        summary=data.get('summary'),

        experience=experience_json,
        education=education_json,
        references=references_json,

        skills=skills_json,
        languages=languages_json,
        interests=interests_json,

        template=template_value
    )

    db.session.add(resume)
    db.session.commit()

    print(f"✅ Saved resume with ID: {resume.id} (template={resume.template})")
    print("📥 SKILLS SENT:", data.get("skills"))
    print("INTERESTS RECEIVED:", data.get("interests"))
    print("LANGUAGES RECEIVED:", data.get("languages"))
    print("REFERENCES RECEIVED:", data.get("references"))

    return jsonify({
        "message": "Resume submitted successfully!",
        "resume_id": resume.id,
        "template": resume.template
    }), 201
set
@app.route("/ai/generate-summary", methods=["POST"])
@login_required
def generate_summary():

    data = request.get_json()
    job_title = data.get("job_title")

    if not job_title:
        return jsonify({"error": "Job title required"}), 400

    prompt = f"""
    Write a professional resume summary for a '{job_title}'.
    Write it as an essay-style professional paragraph suitable for a resume.
    """

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional resume writer."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.6
        )

        summary = response.choices[0].message.content

        return jsonify({"summary": summary})

    except Exception as e:
        print("AI ERROR:", str(e))
        return jsonify({"error": str(e)}), 500

    @app.route("/ai/generate-skills", methods=["POST"])
    def generate_skills():
        data = request.get_json()
        print("Data received:", data)
        job_title = data.get("job_title", "")
        return jsonify({"skills": ["Communication", "Teamwork", "Problem Solving"]})

    @app.route("/ai/generate-skills", methods=["POST"])
    @login_required
    def generate_skills():

        data = request.get_json()
        job_title = data.get("job_title")

        if not job_title:
            return jsonify({"error": "Job title required"}), 400

        prompt = f"""
        Suggest professional resume skills for a '{job_title}'.
        Provide your response as bullet points.
        """

        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a professional resume writer and career coach."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.4
            )

            skills = response.choices[0].message.content
            return jsonify({"skills": skills})

        except Exception as e:
            print("AI ERROR:", str(e))
            return jsonify({"error": str(e)}), 500


# ✅ AI Suggestions Route (Mock)
@app.route('/ai_suggestions', methods=['POST'])
def ai_suggestions():
    data = request.get_json()
    job_title = data.get('job_title')

    if not job_title:
        return jsonify({'error': 'Job title is required for AI suggestions'}), 400

    # Mock suggestions for testing
    mock_suggestions = {
        "summary": f"Enthusiastic {job_title} with a passion for delivering results and driving innovation.",
        "skills": ["Skill 1", "Skill 2", "Skill 3", "Skill 4", "Skill 5"],
        "hobbies": ["Hobby 1", "Hobby 2", "Hobby 3"],
        "experience": f"Worked as a {job_title}, managing tasks and collaborating with cross-functional teams to achieve goals."
    }

    return jsonify({"suggestions": mock_suggestions}), 200

# ✅ Configure Upload Folder
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# ✅ Route to Upload and Parse Resume
@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected for uploading'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    extracted_text = ""
    try:
        if filename.lower().endswith('.pdf'):
            extracted_text = extract_text(filepath)
        elif filename.lower().endswith('.docx'):
            doc = docx.Document(filepath)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
        else:
            return jsonify({'error': 'Unsupported file type. Please upload a PDF or DOCX file.'}), 400
    except Exception as e:
        return jsonify({'error': f'Error parsing file: {str(e)}'}), 500

    # For simplicity, return raw extracted text for frontend review
    return jsonify({'extracted_text': extracted_text}), 200

@app.route('/resume-scanner')
def resume_scanner():
    return render_template("resume_scanner.html")





# app.py (Make sure to add this import at the top of the file)
from ai_engine import match_score_openai
from flask import request, jsonify


# app.py (Replace the existing /scan_resume route)

@app.route("/scan_resume", methods=["POST"])
def scan_resume():
    # Use a global try/except to ensure we ALWAYS return JSON on failure
    try:
        # 1. Get files and text from the request
        resume_file = request.files.get("resume_file")
        jd_file = request.files.get("jd_file")
        jd_text = request.form.get("job_description_text")

        if not resume_file:
            return jsonify({"error": "Resume file is missing."}), 400

        # --- CONFIG CHECK ---
        upload_folder = app.config.get('UPLOAD_FOLDER', 'uploads')
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)
        # --- END CONFIG CHECK ---

        # --- A. EXTRACT RESUME TEXT ---
        resume_text = ""
        try:
            resume_filename = secure_filename(resume_file.filename)
            resume_filepath = os.path.join(upload_folder, resume_filename)
            resume_file.save(resume_filepath)
            resume_ext = resume_filename.rsplit('.', 1)[-1].lower()

            if resume_ext == 'pdf':
                resume_text = extract_text_from_pdf(resume_filepath)
            elif resume_ext in ('docx', 'doc'):
                resume_text = extract_text_from_docx(resume_filepath)

            os.remove(resume_filepath)  # Clean up temporary file
        except Exception as e:
            return jsonify({"error": f"Failed to process resume file: {type(e).__name__}: {str(e)}"}), 500

        if "Error" in resume_text or not resume_text:
            return jsonify({"error": f"Failed to extract readable text from resume. Check file format."}), 500

        # --- B. EXTRACT JOB DESCRIPTION TEXT ---
        job_description = jd_text if jd_text else ""

        if jd_file:
            try:
                jd_filename = secure_filename(jd_file.filename)
                jd_filepath = os.path.join(upload_folder, jd_filename)
                jd_file.save(jd_filepath)
                jd_ext = jd_filename.rsplit('.', 1)[-1].lower()

                if jd_ext == 'pdf':
                    job_description = extract_text_from_pdf(jd_filepath)
                elif jd_ext in ('docx', 'doc'):
                    job_description = extract_text_from_docx(jd_filepath)

                os.remove(jd_filepath)  # Clean up temporary file
            except Exception as e:
                return jsonify({"error": f"Failed to process JD file: {type(e).__name__}: {str(e)}"}), 500

        if not job_description:
            return jsonify({"error": "Job description content is missing. Provide text or a valid file."}), 400

        # 4. Call the OpenAI function
        result = match_score_openai(resume_text, job_description)

        if result and 'error' in result:
            # Catches AI engine errors (e.g., API Key missing)
            return jsonify(result), 500
        elif result:
            # Success
            return jsonify(result)
        else:
            return jsonify({"error": "AI Engine returned null analysis."}), 500

    except Exception as e:
        # GLOBAL CATCH: Catches any unhandled exception (like a missing import/config)
        # This will now return JSON instead of HTML
        print(f"CRITICAL UNHANDLED ERROR in /scan_resume: {type(e).__name__}: {str(e)}")
        return jsonify({
        "error": f"A critical server error occurred. Check the Flask terminal for '{type(e).__name__}' details."}), 500

@app.route('/upload_job_description', methods=['POST'])
def upload_job_description():
    if 'job_description_file' not in request.files:
        return jsonify({"error": "No job description file uploaded"}), 400

    jd_file = request.files['job_description_file']

    if jd_file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    # Save uploaded JD file
    upload_folder = "uploads"
    os.makedirs(upload_folder, exist_ok=True)
    file_path = os.path.join(upload_folder, jd_file.filename)
    jd_file.save(file_path)

    # Extract text
    jd_text = extract_text_from_file(file_path)

    if not jd_text:
        return jsonify({"error": "Could not extract text from JD file"}), 400

    return jsonify({"job_description_text": jd_text})

from flask import send_file

@app.route("/resume_1/<template_id>")
def get_resume_1(template_id):
    file_path = TEMPLATE_FILES.get(template_id)
    if not file_path:
        return "Template not found", 404
    return send_file(file_path, as_attachment=False)



@app.route("/resume_template/<int:resume_id>")
def resume_template(resume_id):
    resume = Resume.query.get_or_404(resume_id)

    import json
    def safe_load(raw, fallback):
        try:
            return json.loads(raw) if raw else fallback
        except:
            return fallback

    experience_list = safe_load(resume.experience, [])
    education_list = safe_load(resume.education, [])
    references_list = safe_load(resume.references, [])

    skills_list = [s.strip() for s in (resume.skills or "").split(",") if s.strip()]
    languages_list = [l.strip() for l in (resume.languages or "").split(",") if l.strip()]
    interests_list = [i.strip() for i in (resume.interests or "").split(",") if i.strip()]

    # 🔥 Dynamically load chosen template
    template_file = f"resume_{resume.template}.html"

    return render_template(
        template_file,
        resume=resume,
        experience=experience_list,
        education=education_list,
        references=references_list,
        skills=skills_list,
        languages=languages_list,
        interests=interests_list
    )

@login_required
@app.route('/set_template/<int:resume_id>', methods=['POST'])
def set_template(resume_id):
    data = request.get_json() or {}
    template = data.get("template") or "template1"
    resume = Resume.query.filter_by(id=resume_id, user_id=session["user_id"]).first_or_404()
    resume.template = template
    db.session.commit()
    return jsonify({"status":"ok", "template": resume.template})




# ------------- Preview & Template rendering -------------
@app.route("/preview_resume/<int:resume_id>")
def preview_resume(resume_id):
    resume = Resume.query.filter_by(
        id=resume_id,
        user_id=session["user_id"]
    ).first_or_404()

    import json

    def safe_json(field):
        """Guarantees list output even if corrupted"""
        try:
            return json.loads(field) if field else []
        except:
            return []

    # ✅ FORCE JSON ARRAYS
    experience = safe_json(resume.experience)
    education = safe_json(resume.education)
    references = safe_json(resume.references)

    # ✅ FORCE comma-lists into real arrays
    def safe_json(field):
        try:
            return json.loads(field) if field else []
        except:
            return []

    skills = safe_json(resume.skills)
    languages = safe_json(resume.languages)
    interests = safe_json(resume.interests)
    references = safe_json(resume.references)
    experience = safe_json(resume.experience)
    education = safe_json(resume.education)

    # ✅ DEBUG OUTPUT (TEMP)
    print("PREVIEW DATA ===>")
    print("INTERESTS:", interests)
    print("REFERENCES:", references)
    print("LANGUAGES:", languages)

    # ✅ Choose template dynamically
    template_file = f"resume_{resume.template}.html"

    print("PREVIEW TEMPLATE:", template_file)
    return render_template(
        template_file,
        resume=resume,
        cleaned={
            "skills": skills,
            "languages": languages,
            "interests": interests,
            "education": education,
            "experience": experience,
            "references": references
        },
        skills=skills,
        languages=languages,
        interests=interests,
        education=education,
        experience=experience,
        references=references
    )


# ------------- Direct Download Route -------------
from flask import Response

@app.route("/download_resume/<int:resume_id>")
def download_resume(resume_id):
    resume = Resume.query.get_or_404(resume_id)

    import json

    def safe_load(raw, fallback):
        try:
            return json.loads(raw) if raw else fallback
        except:
            return fallback

    # Decode structured fields
    experience_list = safe_load(resume.experience, [])
    education_list = safe_load(resume.education, [])
    references_list = safe_load(resume.references, [])

    # Comma-separated → list
    skills_list = [s.strip() for s in (resume.skills or "").split(",") if s.strip()]
    languages_list = [l.strip() for l in (resume.languages or "").split(",") if l.strip()]
    interests_list = [i.strip() for i in (resume.interests or "").split(",") if i.strip()]

    # Render full HTML
    html = render_template(
        "resume_template1.html",
        resume=resume,
        experience=experience_list,
        education=education_list,
        references=references_list,
        skills=skills_list,
        languages=languages_list,
        interests=interests_list,
        no_icons=True   # prevents icon CSS from breaking PDF
    )

    # Convert to PDF
    from weasyprint import HTML
    pdf = HTML(string=html).write_pdf()

    return Response(
        pdf,
        mimetype="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=resume_{resume_id}.pdf"
        }
    )


@app.route('/generate_template', methods=['POST'])
def generate_template():
    """
    Generate a PDF from a chosen resume template and send to client.
    """
    data = request.get_json() or {}
    resume_id = data.get('resume_id')
    template = data.get('template', 'template1')

    resume = Resume.query.get_or_404(resume_id)
    cleaned = clean_resume_data(resume)
    template_file = f"resume_{template}.html"

    rendered_html = render_template(template_file, resume=resume, cleaned=cleaned)

    if WEASYPRINT_AVAILABLE:
        try:
            pdf = HTML(string=rendered_html).write_pdf()
            response = make_response(pdf)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename=resume_{resume_id}.pdf'
            return response
        except Exception as e:
            return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500
    else:
        return jsonify({
            "error": "WeasyPrint not available. Install it to enable PDF downloads.",
            "html_preview": rendered_html
        }), 501

    @app.route('/save_resume_pdf/<int:resume_id>', methods=['GET'])
    def save_resume_pdf(resume_id):
        resume = Resume.query.get_or_404(resume_id)
        cleaned = clean_resume_data(resume)

        rendered_html = render_template("resume_template1.html", resume=resume, cleaned=cleaned)
        filename = f"{PDF_FOLDER}/resume_{resume_id}.pdf"

        if WEASYPRINT_AVAILABLE:
            try:
                HTML(string=rendered_html).write_pdf(filename)
                return jsonify({"saved": True, "file": filename})
            except Exception as e:
                return jsonify({"error": f"PDF save failed: {str(e)}"}), 500
        else:
            return jsonify({"error": "WeasyPrint is not installed."}), 500

@app.route("/test")
def test_route():
    return "OK"


@app.route("/")
def home():
    return render_template("resume_builder_form.html")



if __name__ == '__main__':
    app.run(debug=True)


